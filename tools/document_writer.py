from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

from models.resume import Resume


def write_docx(resume: Resume, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # --- Contact header ---
    contact = resume.contact
    name = contact.get("name", "")
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    contact_parts = [v for k, v in contact.items() if k != "name" and v]
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def section_heading(title: str):
        h = doc.add_heading(title, level=1)
        h.runs[0].font.size = Pt(12)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # --- Summary ---
    if resume.summary:
        section_heading("Summary")
        doc.add_paragraph(resume.summary)

    # --- Skills ---
    if resume.skills:
        section_heading("Skills")
        doc.add_paragraph(" • ".join(resume.skills))

    # --- Experience ---
    if resume.experience:
        section_heading("Experience")
        for role in resume.experience:
            p = doc.add_paragraph()
            run = p.add_run(f"{role.title} — {role.company}")
            run.bold = True
            p.add_run(f"  {role.dates}").italic = True
            for bullet in role.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    # --- Education ---
    if resume.education:
        section_heading("Education")
        for edu in resume.education:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")
            dates = edu.get("dates", "")
            gpa = edu.get("gpa", "")
            line = f"{degree} in {field} — {institution}  {dates}"
            if gpa:
                line += f"  GPA: {gpa}"
            doc.add_paragraph(line)

    # --- Extras ---
    for extra in resume.extras:
        category = extra.get("category", "")
        items = extra.get("items", [])
        if category and items:
            section_heading(category)
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

    doc.save(str(output_path))


def write_pdf(resume: Resume, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    dark = colors.HexColor("#1a1a2e")

    name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=20, textColor=dark, spaceAfter=4)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9, alignment=1, spaceAfter=8)
    section_style = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=12, textColor=dark, spaceBefore=10, spaceAfter=4, borderPad=0)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=4)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, leftIndent=12, spaceAfter=2, bulletIndent=0)
    bold_style = ParagraphStyle("Bold", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold")

    contact = resume.contact
    name = contact.get("name", "")
    contact_parts = [v for k, v in contact.items() if k != "name" and v]

    story = [
        Paragraph(name, name_style),
        Paragraph(" &nbsp;|&nbsp; ".join(contact_parts), contact_style),
        HRFlowable(width="100%", thickness=1, color=dark),
        Spacer(1, 6),
    ]

    def add_section(title: str):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Spacer(1, 4))

    if resume.summary:
        add_section("Summary")
        story.append(Paragraph(resume.summary, body_style))

    if resume.skills:
        add_section("Skills")
        story.append(Paragraph(" &nbsp;•&nbsp; ".join(resume.skills), body_style))

    if resume.experience:
        add_section("Experience")
        for role in resume.experience:
            story.append(Paragraph(f"<b>{role.title} — {role.company}</b> &nbsp;&nbsp;<i>{role.dates}</i>", bold_style))
            for bullet in role.bullets:
                story.append(Paragraph(f"• {bullet}", bullet_style))
            story.append(Spacer(1, 4))

    if resume.education:
        add_section("Education")
        for edu in resume.education:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")
            dates = edu.get("dates", "")
            gpa = edu.get("gpa", "")
            line = f"<b>{degree} in {field}</b> — {institution} &nbsp;&nbsp;<i>{dates}</i>"
            if gpa:
                line += f" &nbsp; GPA: {gpa}"
            story.append(Paragraph(line, body_style))

    for extra in resume.extras:
        category = extra.get("category", "")
        items = extra.get("items", [])
        if category and items:
            add_section(category)
            for item in items:
                story.append(Paragraph(f"• {item}", bullet_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)
